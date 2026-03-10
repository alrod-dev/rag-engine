"""Document ingestion pipeline."""

import os
import uuid
from typing import List, Dict, Any, BinaryIO
from datetime import datetime
import csv

from pypdf import PdfReader
from docx import Document as DocxDocument

from backend.models.schemas import TextChunk, DocumentMetadata
from backend.services.chunking import get_chunker
from backend.utils.text_processing import clean_text, extract_metadata, count_tokens


class DocumentIngestionPipeline:
    """Handles document ingestion and processing."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 100, chunking_strategy: str = "recursive"):
        """
        Initialize ingestion pipeline.

        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            chunking_strategy: Chunking strategy to use
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.chunker = get_chunker(chunking_strategy, chunk_size, chunk_overlap)

    def ingest_file(self, file: BinaryIO, filename: str) -> tuple[str, List[TextChunk], DocumentMetadata]:
        """
        Ingest a file and return document ID and chunks.

        Args:
            file: File object to ingest
            filename: Name of the file

        Returns:
            Tuple of (document_id, chunks, metadata)
        """
        file_ext = os.path.splitext(filename)[1].lower()

        # Extract text based on file type
        if file_ext == ".pdf":
            text = self._extract_from_pdf(file)
        elif file_ext == ".docx":
            text = self._extract_from_docx(file)
        elif file_ext == ".txt":
            text = self._extract_from_txt(file)
        elif file_ext == ".csv":
            text = self._extract_from_csv(file)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Clean text
        text = clean_text(text)

        # Generate document ID
        document_id = str(uuid.uuid4())

        # Extract metadata
        file_size = file.seek(0, 2)  # Seek to end
        file.seek(0)  # Reset to beginning
        metadata_dict = extract_metadata(text, filename)

        metadata = DocumentMetadata(
            filename=filename,
            file_type=file_ext[1:],  # Remove the dot
            file_size=file_size,
            upload_date=datetime.utcnow(),
            num_pages=metadata_dict.get("num_pages", None),
            custom_metadata=metadata_dict,
        )

        # Chunk the document
        chunks = self.chunker.chunk(text, document_id, {"source": filename})

        metadata.num_chunks = len(chunks)

        return document_id, chunks, metadata

    def _extract_from_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF."""
        try:
            reader = PdfReader(file)
            text = ""
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            return text
        except Exception as e:
            raise ValueError(f"Error extracting from PDF: {str(e)}")

    def _extract_from_docx(self, file: BinaryIO) -> str:
        """Extract text from DOCX."""
        try:
            doc = DocxDocument(file)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"

            return text
        except Exception as e:
            raise ValueError(f"Error extracting from DOCX: {str(e)}")

    def _extract_from_txt(self, file: BinaryIO) -> str:
        """Extract text from TXT."""
        try:
            content = file.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="replace")
            return content
        except Exception as e:
            raise ValueError(f"Error extracting from TXT: {str(e)}")

    def _extract_from_csv(self, file: BinaryIO) -> str:
        """Extract text from CSV."""
        try:
            text_content = file.read()
            if isinstance(text_content, bytes):
                text_content = text_content.decode("utf-8", errors="replace")

            # Parse CSV and convert to readable format
            lines = text_content.strip().split("\n")
            reader = csv.reader(lines)

            text = ""
            for row_num, row in enumerate(reader):
                text += f"Row {row_num + 1}: {' | '.join(row)}\n"

            return text
        except Exception as e:
            raise ValueError(f"Error extracting from CSV: {str(e)}")

    def update_chunking_config(self, chunk_size: int, chunk_overlap: int, strategy: str):
        """Update chunking configuration."""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = strategy
        self.chunker = get_chunker(strategy, chunk_size, chunk_overlap)

    def get_chunking_config(self) -> Dict[str, Any]:
        """Get current chunking configuration."""
        return {
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "strategy": self.chunking_strategy,
        }
